# import json
# import csv
# import logging
# from typing import List, Dict, Optional
# import tempfile
# import os
# from pathlib import Path

# # PDF processing
# try:
#     import PyPDF2
#     PDF_AVAILABLE = True
# except ImportError:
#     try:
#         import pdfplumber
#         PDF_AVAILABLE = True
#         USE_PDFPLUMBER = True
#     except ImportError:
#         PDF_AVAILABLE = False
#         USE_PDFPLUMBER = False

# # DOCX processing
# try:
#     from docx import Document
#     DOCX_AVAILABLE = True
# except ImportError:
#     DOCX_AVAILABLE = False

# logger = logging.getLogger(__name__)

# class DocumentProcessor:
#     """
#     Processes various document formats and converts them into searchable chunks.
#     """
    
#     def __init__(self, device: str = "cpu", chunk_size: int = 1000, chunk_overlap: int = 200):
#         self.device = device
#         self.chunk_size = chunk_size
#         self.chunk_overlap = chunk_overlap
#         logger.info(f"DocumentProcessor initialized with device: {device}")
    
#     def process_file(self, file_path: str, filename: str) -> List[Dict]:
#         """
#         Process a single file and return document chunks.
        
#         Args:
#             file_path: Path to the file
#             filename: Original filename
            
#         Returns:
#             List of document dictionaries with text and metadata
#         """
#         file_extension = Path(filename).suffix.lower()
        
#         try:
#             if file_extension == '.pdf':
#                 return self._process_pdf(file_path, filename)
#             elif file_extension == '.docx':
#                 return self._process_docx(file_path, filename)
#             elif file_extension == '.txt':
#                 return self._process_txt(file_path, filename)
#             elif file_extension == '.json':
#                 return self._process_json(file_path, filename)
#             elif file_extension == '.csv':
#                 return self._process_csv(file_path, filename)
#             else:
#                 raise ValueError(f"Unsupported file format: {file_extension}")
                
#         except Exception as e:
#             logger.error(f"Error processing {filename}: {e}")
#             raise
    
#     def _process_pdf(self, file_path: str, filename: str) -> List[Dict]:
#         """Process PDF files."""
#         if not PDF_AVAILABLE:
#             raise ImportError("PDF processing libraries not available. Install PyPDF2 or pdfplumber.")
        
#         text_content = ""
        
#         if 'USE_PDFPLUMBER' in globals() and USE_PDFPLUMBER:
#             import pdfplumber
#             with pdfplumber.open(file_path) as pdf:
#                 for page_num, page in enumerate(pdf.pages, 1):
#                     page_text = page.extract_text()
#                     if page_text:
#                         text_content += f"\n[Page {page_num}]\n{page_text}\n"
#         else:
#             with open(file_path, 'rb') as file:
#                 pdf_reader = PyPDF2.PdfReader(file)
#                 for page_num, page in enumerate(pdf_reader.pages, 1):
#                     page_text = page.extract_text()
#                     if page_text:
#                         text_content += f"\n[Page {page_num}]\n{page_text}\n"
        
#         return self._chunk_text(text_content, filename, "pdf")
    
#     def _process_docx(self, file_path: str, filename: str) -> List[Dict]:
#         """Process DOCX files."""
#         if not DOCX_AVAILABLE:
#             raise ImportError("DOCX processing library not available. Install python-docx.")
        
#         doc = Document(file_path)
#         text_content = ""
        
#         for paragraph in doc.paragraphs:
#             if paragraph.text.strip():
#                 text_content += paragraph.text + "\n"
        
#         # Process tables if any
#         for table in doc.tables:
#             for row in table.rows:
#                 row_text = " | ".join(cell.text.strip() for cell in row.cells)
#                 if row_text.strip():
#                     text_content += row_text + "\n"
        
#         return self._chunk_text(text_content, filename, "docx")
    
#     def _process_txt(self, file_path: str, filename: str) -> List[Dict]:
#         """Process TXT files."""
#         try:
#             with open(file_path, 'r', encoding='utf-8') as file:
#                 text_content = file.read()
#         except UnicodeDecodeError:
#             # Try with different encoding
#             with open(file_path, 'r', encoding='latin-1') as file:
#                 text_content = file.read()
        
#         return self._chunk_text(text_content, filename, "txt")
    
#     def _process_json(self, file_path: str, filename: str) -> List[Dict]:
#         """Process JSON files."""
#         with open(file_path, 'r', encoding='utf-8') as file:
#             data = json.load(file)
        
#         documents = []
        
#         if isinstance(data, list):
#             for i, item in enumerate(data):
#                 if isinstance(item, dict):
#                     # If item has 'text' and 'id' fields, use them directly
#                     if 'text' in item and 'id' in item:
#                         documents.append({
#                             'id': item['id'],
#                             'text': item['text'],
#                             'metadata': {
#                                 'filename': filename,
#                                 'file_type': 'json',
#                                 'item_index': i
#                             }
#                         })
#                     else:
#                         # Convert entire item to text
#                         text_content = json.dumps(item, indent=2)
#                         documents.append({
#                             'id': f"{filename}_{i}",
#                             'text': text_content,
#                             'metadata': {
#                                 'filename': filename,
#                                 'file_type': 'json',
#                                 'item_index': i
#                             }
#                         })
#                 else:
#                     # Handle non-dict items
#                     documents.append({
#                         'id': f"{filename}_{i}",
#                         'text': str(item),
#                         'metadata': {
#                             'filename': filename,
#                             'file_type': 'json',
#                             'item_index': i
#                         }
#                     })
#         else:
#             # Single object
#             if isinstance(data, dict) and 'text' in data and 'id' in data:
#                 documents.append({
#                     'id': data['id'],
#                     'text': data['text'],
#                     'metadata': {
#                         'filename': filename,
#                         'file_type': 'json'
#                     }
#                 })
#             else:
#                 text_content = json.dumps(data, indent=2)
#                 documents.append({
#                     'id': filename,
#                     'text': text_content,
#                     'metadata': {
#                         'filename': filename,
#                         'file_type': 'json'
#                     }
#                 })
        
#         return documents
    
#     def _process_csv(self, file_path: str, filename: str) -> List[Dict]:
#         """Process CSV files."""
#         text_content = ""
        
#         with open(file_path, 'r', encoding='utf-8') as file:
#             csv_reader = csv.reader(file)
#             rows = list(csv_reader)
            
#             if rows:
#                 # First row as headers
#                 headers = rows[0]
#                 text_content += f"Headers: {', '.join(headers)}\n\n"
                
#                 # Process each row
#                 for i, row in enumerate(rows[1:], 1):
#                     row_text = ", ".join(str(cell) for cell in row)
#                     text_content += f"Row {i}: {row_text}\n"
        
#         return self._chunk_text(text_content, filename, "csv")
    
#     def _chunk_text(self, text: str, filename: str, file_type: str) -> List[Dict]:
#         """
#         Split text into chunks for better search performance.
        
#         Args:
#             text: The text to chunk
#             filename: Original filename
#             file_type: Type of the file
            
#         Returns:
#             List of document chunks
#         """
#         if not text.strip():
#             return []
        
#         # Simple chunking strategy
#         chunks = []
#         words = text.split()
        
#         for i in range(0, len(words), self.chunk_size - self.chunk_overlap):
#             chunk_words = words[i:i + self.chunk_size]
#             chunk_text = " ".join(chunk_words)
            
#             if chunk_text.strip():
#                 chunks.append({
#                     'id': f"{filename}_chunk_{len(chunks)}",
#                     'text': chunk_text,
#                     'metadata': {
#                         'filename': filename,
#                         'file_type': file_type,
#                         'chunk_index': len(chunks),
#                         'char_count': len(chunk_text),
#                         'word_count': len(chunk_words)
#                     }
#                 })
        
#         # If text is short, create single chunk
#         if not chunks:
#             chunks.append({
#                 'id': f"{filename}_chunk_0",
#                 'text': text,
#                 'metadata': {
#                     'filename': filename,
#                     'file_type': file_type,
#                     'chunk_index': 0,
#                     'char_count': len(text),
#                     'word_count': len(text.split())
#                 }
#             })
        
#         logger.info(f"Created {len(chunks)} chunks for {filename}")
#         return chunks
    
#     def get_supported_formats(self) -> List[str]:
#         """Return list of supported file formats."""
#         formats = ['txt', 'json', 'csv']
        
#         if PDF_AVAILABLE:
#             formats.append('pdf')
            
#         if DOCX_AVAILABLE:
#             formats.append('docx')
            
#         return formats
import json
import csv
import logging
from typing import List, Dict, Optional
from pathlib import Path

# PDF processing
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    try:
        import pdfplumber
        PDF_AVAILABLE = True
        USE_PDFPLUMBER = True
    except ImportError:
        PDF_AVAILABLE = False
        USE_PDFPLUMBER = False

# DOCX processing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Import LLMHandler for type hinting, avoiding circular import
from typing import TYPE_CHECKING
if TYPE_CHECKING:
    from .llm_handler import LLMHandler

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self, device: str = "cpu", chunk_size: int = 1000, chunk_overlap: int = 200):
        self.device = device
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        logger.info(f"DocumentProcessor initialized with device: {device}")

    def process_file(self, file_path: str, filename: str, llm_handler: Optional['LLMHandler'] = None) -> List[Dict]:
        file_extension = Path(filename).suffix.lower()

        try:
            if file_extension == '.json':
                return self._process_json(file_path, filename)

            text_content = ""
            if file_extension == '.pdf':
                text_content = self._extract_text_from_pdf(file_path)
            elif file_extension == '.docx':
                text_content = self._extract_text_from_docx(file_path)
            elif file_extension == '.txt':
                text_content = self._extract_text_from_txt(file_path)
            elif file_extension == '.csv':
                text_content = self._extract_text_from_csv(file_path, filename)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")

            if llm_handler and text_content.strip():
                logger.info(f"Structuring {filename} using LLM...")
                return self._structure_text_with_llm(text_content, filename, file_extension[1:], llm_handler)
            elif text_content.strip():
                logger.warning(f"No LLM handler provided or content is empty. Falling back to simple chunking for {filename}.")
                return self._chunk_text(text_content, filename, file_extension[1:])
            else:
                return []

        except Exception as e:
            logger.error(f"Error processing {filename}: {e}")
            raise

    def _structure_text_with_llm(self, text_content: str, filename: str, file_type: str, llm_handler: 'LLMHandler') -> List[Dict]:
        prompt = f"""
You are an intelligent assistant trained to extract structured information from documents for a knowledge retrieval system (RAG).

Given a raw document text, extract the most important **Question: ... Answer: ...** style information that can help someone query this document later. Each output must be in the following JSON format:

```json
{{
  "text": "Question: <question> Answer: <answer>",
  "metadata": {{
    "source": "{filename}",
    "category": "<category>",
    "keywords": ["<keyword1>", "<keyword2>", "<keyword3>"]
  }}
}}
```

**Instructions:**
1. **Extract Q&A pairs** from the document based on distinct ideas or topics.
2. **Generate clear and concise questions and answers** that reflect the most useful information.
3. **Category** must be one of: ["Overview", "Process", "Details", "Examples", "Concepts", "Issues", "Notes"]
4. **Keywords** must be 3–7 relevant terms per entry.
5. **Source** must always be "{filename}".
6. **Return** only a **valid JSON array**. Do not include any markdown, explanation, or extra text.

**Document Content:**
{text_content}
"""

        try:
            response_text = llm_handler.generate_response(prompt)
            json_str = self._extract_json_from_response(response_text)
            try:
                structured_data = json.loads(json_str)
                if isinstance(structured_data, dict):
                    structured_data = [structured_data]
            except json.JSONDecodeError:
                structured_data = self._recover_json_objects(json_str, filename)

            final_docs = []
            for i, item in enumerate(structured_data):
                if isinstance(item, dict):
                    item['id'] = f"{filename}_{i}"
                    item.setdefault('metadata', {})
                    item['metadata']['source'] = filename
                    item['metadata']['file_type'] = file_type
                    final_docs.append(item)

            if not final_docs:
                raise ValueError("No valid JSON objects recovered.")
            return final_docs

        except Exception as e:
            logger.error(f"Failed to parse or recover LLM response for {filename}. Error: {e}")
            return self._chunk_text(text_content, filename, file_type)

    def _extract_json_from_response(self, response_text: str) -> str:
        import re
        code_block_match = re.search(r"```(?:json)?\s*(\[[\s\S]*?\])\s*```", response_text)
        if code_block_match:
            return code_block_match.group(1)
        array_match = re.search(r"(\[[\s\S]*\])", response_text)
        if array_match:
            return array_match.group(1)
        object_match = re.search(r"(\{[\s\S]*\})", response_text)
        if object_match:
            return f"[{object_match.group(1)}]"
        return response_text.strip()

    def _recover_json_objects(self, json_str: str, filename: str) -> List[Dict]:
        import re
        recovered_objects = []
        brace_count = 0
        start_pos = -1
        for i, char in enumerate(json_str):
            if char == '{':
                if brace_count == 0:
                    start_pos = i
                brace_count += 1
            elif char == '}':
                brace_count -= 1
                if brace_count == 0 and start_pos != -1:
                    obj_str = json_str[start_pos:i+1]
                    try:
                        cleaned_obj = re.sub(r',\s*}', '}', obj_str)
                        obj = json.loads(cleaned_obj)
                        recovered_objects.append(obj)
                    except json.JSONDecodeError:
                        continue
                    start_pos = -1
        return recovered_objects

    def _extract_text_from_pdf(self, file_path: str) -> str:
        if not PDF_AVAILABLE:
            raise ImportError("PDF processing libraries not available.")
        text_content = ""
        if 'USE_PDFPLUMBER' in globals() and USE_PDFPLUMBER:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_content += text + "\n"
        else:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_content += text + "\n"
        return text_content

    def _extract_text_from_docx(self, file_path: str) -> str:
        if not DOCX_AVAILABLE:
            raise ImportError("DOCX processing library not available.")
        doc = Document(file_path)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                text += "\n" + " | ".join(cell.text.strip() for cell in row.cells)
        return text

    def _extract_text_from_txt(self, file_path: str) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()

    def _extract_text_from_csv(self, file_path: str, filename: str) -> str:
        text = f"Summary of CSV file: {filename}\n"
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            reader = csv.reader(f)
            rows = list(reader)
            if rows:
                headers = rows[0]
                text += f"Headers: {', '.join(headers)}\n\n"
                for i, row in enumerate(rows[1:], 1):
                    row_text = ", ".join(f"{headers[j]}: {cell}" for j, cell in enumerate(row) if j < len(headers))
                    text += f"Row {i}: {row_text}\n"
        return text

    def _process_json(self, file_path: str, filename: str) -> List[Dict]:
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        if isinstance(data, list):
            for i, doc in enumerate(data):
                doc['id'] = f"{filename}_{i}"
            return data
        elif isinstance(data, dict):
            data['id'] = f"{filename}_0"
            return [data]
        return []

    def _chunk_text(self, text: str, filename: str, file_type: str) -> List[Dict]:
        if not text.strip():
            return []
        chunks = []
        words = text.split()
        for i in range(0, len(words), self.chunk_size - self.chunk_overlap):
            chunk_words = words[i:i + self.chunk_size]
            chunk_text = " ".join(chunk_words)
            chunks.append({
                'id': f"{filename}_chunk_{len(chunks)}",
                'text': chunk_text,
                'metadata': {
                    'filename': filename,
                    'file_type': file_type,
                    'chunk_index': len(chunks),
                    'char_count': len(chunk_text),
                    'word_count': len(chunk_words)
                }
            })
        return chunks

    def get_supported_formats(self) -> List[str]:
        formats = ['txt', 'json', 'csv']
        if PDF_AVAILABLE:
            formats.append('pdf')
        if DOCX_AVAILABLE:
            formats.append('docx')
        return formats
