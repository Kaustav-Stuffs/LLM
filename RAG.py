import os
import json
import docx
import pandas as pd
import uvicorn
import time
from fastapi import FastAPI, File, UploadFile, HTTPException, Request
from fastapi.responses import JSONResponse
from pydantic import BaseModel
from langchain_community.llms import Ollama
from langchain_ollama import ChatOllama
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.vectorstores import Chroma
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.messages import HumanMessage, AIMessage
from rank_bm25 import BM25Okapi
from typing import List, Dict, Optional
from contextlib import asynccontextmanager

class QuestionRequest(BaseModel):
    question: str
    chat_history: Optional[List[Dict[str, str]]] = None

class DocumentQA:
    def __init__(self, model_name="llama3.2:3b", embedding_model="all-MiniLM-L6-v2"):
        """
        Initialize the Document QA system with Llama 3.2 model and Hugging Face embeddings
        """
        # Initialize Ollama LLM
        self.llm = ChatOllama(model=model_name, temperature=0.3)
        
        # Initialize HuggingFace embeddings for vector search
        try:
            self.embeddings = HuggingFaceEmbeddings(model_name=embedding_model, 
                                                   model_kwargs={"local_files_only": False})
            print(f"Successfully loaded embedding model: {embedding_model}")
        except Exception as e:
            print(f"[ERROR] Could not load embedding model: {e}")
            print("If you see an 'Unauthorized' error, try running 'huggingface-cli login' or remove any invalid Hugging Face tokens from your environment.")
            raise
        
        # Vector stores
        self.vectorstore = None
        self.bm25 = None
        
        # Document storage
        self.documents = []
        self.chunks = []
        
        # Chat history management
        self.chat_history = []
        self.max_history_length = 5
        
        # Text splitter for document chunking
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len
        )

    def read_document(self, file_path: str) -> List[Document]:
        """
        Read documents from various file formats
        """
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                return [Document(page_content=text, metadata={"source": file_path})]
            
            elif file_ext == '.docx':
                doc = docx.Document(file_path)
                text = '\n'.join([para.text for para in doc.paragraphs])
                return [Document(page_content=text, metadata={"source": file_path})]
            
            elif file_ext == '.csv':
                df = pd.read_csv(file_path)
                text = df.to_string()
                return [Document(page_content=text, metadata={"source": file_path})]
            
            elif file_ext == '.json':
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                text = json.dumps(data, indent=2)
                return [Document(page_content=text, metadata={"source": file_path})]
            
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
        
        except Exception as e:
            print(f"Error reading {file_path}: {e}")
            return []

    def process_documents(self, document_paths: List[str]):
        """
        Process multiple documents and create vector store
        """
        self.documents = []
        for path in document_paths:
            self.documents.extend(self.read_document(path))
        
        # Split documents into chunks
        print(f"Splitting {len(self.documents)} documents into chunks...")
        all_chunks = self.text_splitter.split_documents(self.documents)
        self.chunks = [chunk.page_content for chunk in all_chunks]
        chunk_metadata = [chunk.metadata for chunk in all_chunks]
        
        print(f"Created {len(self.chunks)} chunks for indexing")
        
        # Create FAISS vector store
        print("Initializing FAISS vector store with embeddings...")
        self.vectorstore = FAISS.from_texts(self.chunks, self.embeddings, metadatas=chunk_metadata)
        
        # Initialize BM25 for keyword search
        print("Initializing BM25 for keyword search...")
        tokenized_corpus = [chunk.split() for chunk in self.chunks]
        self.bm25 = BM25Okapi(tokenized_corpus)
        
        print("Document processing complete. Ready for queries.")

    def hybrid_search(self, query: str, chat_history: List[Dict[str, str]] = None, k: int = 5):
        """
        Perform hybrid search combining semantic (FAISS) and keyword (BM25) search
        """
        if not self.vectorstore or not self.bm25:
            return []
        
        # Expand query with context from chat history
        expanded_query = query
        if chat_history:
            # Extract only relevant context terms from recent history
            recent_messages = chat_history[-self.max_history_length:]
            context_terms = []
            for msg in recent_messages:
                # Extract meaningful terms (skip common words)
                question_terms = [term for term in msg.get('question', '').split() 
                                if len(term) > 3 and term.lower() not in ('what', 'when', 'where', 'how', 'why', 'who', 'is', 'are', 'the', 'this', 'that')]
                context_terms.extend(question_terms)
            
            # Add relevant context terms to query
            if context_terms:
                expanded_query += " " + " ".join(context_terms)
        
        print(f"Expanded query: {expanded_query}")
        
        # Semantic search with FAISS
        vector_results = self.vectorstore.similarity_search_with_score(expanded_query, k=k)
        vector_docs = [(doc.page_content, score, doc.metadata) for doc, score in vector_results]
        
        # BM25 keyword search
        tokenized_query = expanded_query.split()
        bm25_scores = self.bm25.get_scores(tokenized_query)
        bm25_top_k = sorted(enumerate(bm25_scores), key=lambda x: x[1], reverse=True)[:k]
        bm25_docs = [(self.chunks[i], score, {}) for i, score in bm25_top_k if i < len(self.chunks)]
        
        # Combine results (weighted average)
        combined_results = {}
        
        # Add vector search results (lower score is better in FAISS)
        for content, score, metadata in vector_docs:
            # Transform FAISS score (lower is better) to a 0-1 scale where higher is better
            normalized_score = max(0, 1 - score)  # Simple normalization
            combined_results[content] = combined_results.get(content, 0) + normalized_score * 0.6
        
        # Add BM25 results
        for content, score, _ in bm25_docs:
            # Normalize BM25 scores to 0-1 range
            max_bm25 = max([s for _, s in bm25_top_k]) if bm25_top_k else 1
            normalized_score = score / max_bm25 if max_bm25 > 0 else 0
            combined_results[content] = combined_results.get(content, 0) + normalized_score * 0.4
        
        # Sort by combined score and get top results
        sorted_results = sorted(combined_results.items(), key=lambda x: x[1], reverse=True)[:k]
        
        # Convert back to Document format
        result_docs = [Document(page_content=content) for content, _ in sorted_results]
        
        return result_docs

    def answer_question(self, query: str, chat_history: List[Dict[str, str]] = None):
        """
        Answer question using hybrid search and LLM
        """
        if not self.vectorstore:
            return "No documents have been processed yet. Please upload documents first."
        
        start_time = time.time()
        
        # Format chat history for context
        history_str = ""
        if chat_history:
            for msg in chat_history[-self.max_history_length:]:
                q = msg.get('question', '')
                a = msg.get('answer', '')
                if q:
                    history_str += f"Human: {q}\n"
                if a:
                    history_str += f"Assistant: {a}\n"
        
        # Perform hybrid search
        context_docs = self.hybrid_search(query, chat_history)
        
        # If no relevant context found
        if not context_docs:
            return "I couldn't find any relevant information to answer your question. Could you please rephrase or ask something related to the uploaded documents?"
        
        # Prepare context for LLM
        context = "\n\n".join([doc.page_content for doc in context_docs])
        
        # Create prompt template
        prompt_template = ChatPromptTemplate.from_messages([
            ("system", """You are a helpful assistant that answers questions based only on the provided document context. 
            Use the chat history to understand context but only answer based on the document content.
            If the question cannot be answered with the provided documents, respond with:
            "I couldn't find information about that in the documents. Could you please ask something related to the content of the uploaded files?"
            
            Chat History:
            {chat_history}
            
            Document Context:
            {context}
            
            Question: {question}"""),
            ("human", "{question}")
        ])
        
        # Generate response
        chain = prompt_template | self.llm
        try:
            response = chain.invoke({
                "question": query,
                "context": context,
                "chat_history": history_str
            })
            answer = response.content
            
            # Update internal chat history
            self.chat_history.append(HumanMessage(content=query))
            self.chat_history.append(AIMessage(content=answer))
            if len(self.chat_history) > self.max_history_length * 2:
                self.chat_history = self.chat_history[-self.max_history_length * 2:]
                
            print(f"Query answered in {time.time() - start_time:.2f} seconds")
            return answer
        except Exception as e:
            print(f"Error generating response: {e}")
            return "I encountered an error processing your question. Please try again or rephrase your question."

# FastAPI Application
app = FastAPI(title="Llama 3.2 Document QA")
document_qa = DocumentQA()

@app.post("/upload-documents")
async def upload_documents(files: List[UploadFile] = File(...)):
    """
    Upload and process documents
    """
    try:
        # Save uploaded files
        document_paths = []
        for file in files:
            file_path = os.path.join("uploaded_docs", file.filename)
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            
            with open(file_path, "wb") as buffer:
                buffer.write(await file.read())
            
            document_paths.append(file_path)
        
        # Process documents
        document_qa.process_documents(document_paths)
        
        return {"message": "Documents processed successfully"}
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/ask")
async def ask_question(request: QuestionRequest):
    """
    Answer questions based on uploaded documents
    """
    try:
        # Get question and chat history
        question = request.question
        chat_history = request.chat_history or []
        
        # Answer question
        answer = document_qa.answer_question(question, chat_history)
        
        return JSONResponse(content={
            "question": question,
            "answer": answer
        })
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    # Ensure uploaded_docs directory exists
    os.makedirs("uploaded_docs", exist_ok=True)
    
    # Run FastAPI server
    uvicorn.run(app, host="0.0.0.0", port=5556)