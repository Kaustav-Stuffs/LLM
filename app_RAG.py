# # import streamlit as st
# # import os
# # import json
# # import csv
# # import logging
# # from typing import List
# # from threading import Lock

# # from RAG1 import HybridSearchRAG
# # from webcamaccess import process_query
# # import google.generativeai as genai
# # from langchain_ollama import ChatOllama

# # # --- Logging ---
# # logging.basicConfig(level=logging.WARNING)
# # logger = logging.getLogger(__name__)

# # # --- Gemini Setup ---
# # genai.configure(api_key=os.getenv("GEMINI_API_KEY", "AIzaSyDxYogQnjYmTbvghsqSWCKtot5DpygC8hA"))
# # gemini_model = genai.GenerativeModel("gemini-2.5-flash")

# # # --- Qwen2.5 Setup ---
# # qwen_model = ChatOllama(model="phi3:3.8b", temperature=0.1, max_tokens=128)

# # # --- Document Loading ---
# # def load_documents(file_paths: List[str]) -> List[dict]:
# #     documents = []
# #     for path in file_paths:
# #         try:
# #             if path.endswith(".json"):
# #                 with open(path, 'r', encoding='utf-8') as f:
# #                     data = json.load(f)
# #                     if isinstance(data, list):
# #                         documents.extend(data)
# #                     else:
# #                         documents.append(data)
# #             elif path.endswith(".txt"):
# #                 with open(path, 'r', encoding='utf-8') as f:
# #                     documents.append({"text": f.read(), "metadata": {}})
# #             elif path.endswith(".csv"):
# #                 with open(path, 'r', encoding='utf-8') as f:
# #                     reader = csv.reader(f)
# #                     text = "\n".join([", ".join(row) for row in reader])
# #                     documents.append({"text": text, "metadata": {}})
# #         except Exception as e:
# #             logger.error(f"Error reading file {path}: {str(e)}")
# #     return documents

# # DOCUMENT_PATHS = ["./doc1.json"]
# # search_engine_lock = Lock()

# # @st.cache_resource(show_spinner=True)
# # def get_search_engine():
# #     documents = load_documents(DOCUMENT_PATHS)
# #     if not documents:
# #         st.error("No documents loaded. Please check doc1.json.")
# #         return None
# #     return HybridSearchRAG(
# #         documents=documents,
# #         vector_weight=0.7,
# #         bm25_weight=0.3,
# #         top_k=3
# #     )

# # search_engine = get_search_engine()

# # # --- Streamlit UI ---
# # st.set_page_config(page_title="SFA Chatbot", page_icon="🤖", layout="centered")
# # st.title("🤖 SFA Customer Support Chatbot")
# # st.markdown(
# #     """
# #     This chatbot answers questions about the SFA application using RAG and LLMs.<br>
# #     <small>Choose a model, ask your question, and get instant answers!</small>
# #     """,
# #     unsafe_allow_html=True,
# # )

# # if "chat_history" not in st.session_state:
# #     st.session_state.chat_history = []

# # if "user_id" not in st.session_state:
# #     st.session_state.user_id = 1  # For demo, static user id

# # st.sidebar.title("Settings")
# # model_choice = st.sidebar.radio("Choose Model", ("gemini", "llama"), horizontal=True)

# # for entry in st.session_state.chat_history:
# #     if entry["role"] == "user":
# #         st.chat_message("user").write(entry["content"])
# #     else:
# #         st.chat_message("assistant").write(entry["content"])

# # user_input = st.chat_input("Type your question about SFA...")

# # if user_input:
# #     st.session_state.chat_history.append({"role": "user", "content": user_input})
# #     st.chat_message("user").write(user_input)

# #     if not search_engine:
# #         answer = "Sorry, search system is not ready."
# #     else:
# #         try:
# #             result = process_query(
# #                 query=user_input,
# #                 min_relevance_score=0.5,
# #                 model_choice=model_choice,
# #                 gemini_model=gemini_model,
# #                 qwen_model=qwen_model,
# #                 search_engine=search_engine
# #             )
# #             answer = result.get("answer", "No answer received.")
# #         except Exception as e:
# #             answer = f"Error processing your question: {e}"

# #     st.session_state.chat_history.append({"role": "assistant", "content": answer})
# #     st.chat_message("assistant").write(answer)
# import streamlit as st
# import os
# import json
# import csv
# import logging
# from typing import List
# from threading import Lock

# from RAG1 import HybridSearchRAG
# from webcamaccess import process_query
# import google.generativeai as genai
# from langchain_ollama import ChatOllama
# from fileupdater import convert_paragraph_to_json
# import docx
# import PyPDF2
# # --- Logging ---
# logging.basicConfig(level=logging.WARNING)
# logger = logging.getLogger(__name__)

# # --- Gemini Setup ---
# genai.configure(api_key=os.getenv("GEMINI_API_KEY", "AIzaSyDxYogQnjYmTbvghsqSWCKtot5DpygC8hA"))
# gemini_model = genai.GenerativeModel("gemini-2.0-flash")

# # --- Qwen2.5 Setup ---
# qwen_model = ChatOllama(model="phi3:3.8b", temperature=0.1, max_tokens=128)

# # --- Document Loading ---
# def load_documents(file_paths: List[str]) -> List[dict]:
#     documents = []
#     for path in file_paths:
#         try:
#             if path.endswith(".json"):
#                 with open(path, 'r', encoding='utf-8') as f:
#                     data = json.load(f)
#                     if isinstance(data, list):
#                         documents.extend(data)
#                     else:
#                         documents.append(data)
#             elif path.endswith(".txt"):
#                 with open(path, 'r', encoding='utf-8') as f:
#                     documents.append({"text": f.read(), "metadata": {}})
#             elif path.endswith(".csv"):
#                 with open(path, 'r', encoding='utf-8') as f:
#                     reader = csv.reader(f)
#                     text = "\n".join([", ".join(row) for row in reader])
#                     documents.append({"text": text, "metadata": {}})
#         except Exception as e:
#             logger.error(f"Error reading file {path}: {str(e)}")
#     return documents

# # --- Remove default document path ---
# DOCUMENT_PATHS = []  # Start empty, do not load doc1.json

# search_engine_lock = Lock()

# @st.cache_resource(show_spinner=True)
# def get_search_engine():
#     if not DOCUMENT_PATHS:
#         return None
#     documents = load_documents(DOCUMENT_PATHS)
#     if not documents:
#         st.error("No documents loaded. Please upload a knowledge base file.")
#         return None
#     return HybridSearchRAG(
#         documents=documents,
#         vector_weight=0.7,
#         bm25_weight=0.3,
#         top_k=3
#     )

# # Only set search_engine after upload
# if "search_engine" not in st.session_state:
#     st.session_state.search_engine = None

# # --- RAG File Upload Feature ---
# st.sidebar.markdown("---")
# st.sidebar.subheader("Update Knowledge Base")
# uploaded_file = st.sidebar.file_uploader(
#     "Upload new RAG file (.json, .txt, .csv, .pdf, .docx)", 
#     type=["json", "txt", "csv", "pdf", "docx"],
#     help="Replace the current knowledge base with a new file."
# )

# # Add a button to confirm update
# if uploaded_file:
#     st.sidebar.success(f"Uploaded and saved as uploaded_files/{uploaded_file.name}")
#     if st.sidebar.button("Update Knowledge Base"):
#         os.makedirs("uploaded_files", exist_ok=True)
#         raw_path = os.path.join("uploaded_files", uploaded_file.name)
#         with open(raw_path, "wb") as f:
#             f.write(uploaded_file.getbuffer())

#         ext = os.path.splitext(uploaded_file.name)[1].lower()
#         if ext == ".json":
#             json_path = raw_path
#         else:
#             uploaded_file.seek(0)
#             if ext == ".txt":
#                 paragraph = uploaded_file.read().decode("utf-8")
#             elif ext == ".csv":
#                 paragraph = uploaded_file.read().decode("utf-8")
#             elif ext == ".pdf":
#                 try:
#                     pdf_reader = PyPDF2.PdfReader(uploaded_file)
#                     paragraph = ""
#                     for page in pdf_reader.pages:
#                         paragraph += page.extract_text() or ""
#                 except Exception as e:
#                     st.sidebar.error(f"PDF extraction failed: {e}")
#                     paragraph = ""
#             elif ext == ".docx":
#                 try:
#                     doc = docx.Document(uploaded_file)
#                     paragraph = "\n".join([para.text for para in doc.paragraphs])
#                 except Exception as e:
#                     st.sidebar.error(f"DOCX extraction failed: {e}")
#                     paragraph = ""
#             else:
#                 paragraph = ""

#             if paragraph.strip():
#                 new_entries = convert_paragraph_to_json(paragraph)
#                 base_name = os.path.splitext(uploaded_file.name)[0]
#                 json_path = os.path.join("uploaded_files", f"{base_name}.json")
#                 with open(json_path, "w", encoding="utf-8") as jf:
#                     json.dump(new_entries, jf, indent=2, ensure_ascii=False)
#                 st.sidebar.info(f"Converted and saved as {json_path}")
#             else:
#                 st.sidebar.error("No text could be extracted from the uploaded file.")
#                 json_path = None

#         if json_path:
#             DOCUMENT_PATHS.clear()
#             DOCUMENT_PATHS.append(json_path)
#             get_search_engine.clear()
#             st.session_state.search_engine = get_search_engine()
#             st.session_state.chat_history = []
#             st.session_state.last_uploaded = uploaded_file.name  # Track last upload
#             st.sidebar.success("Knowledge base updated! You can now ask questions about the new document.")

# # --- Block Chat Until Knowledge Base is Loaded ---
# def is_kb_loaded():
#     return st.session_state.get("search_engine") is not None

# if not is_kb_loaded():
#     st.warning("Please upload a knowledge base file (.json, .txt, .csv, .pdf, .docx) to start using the chatbot.")
#     st.stop()

# # --- Streamlit UI ---
# st.set_page_config(page_title="RAG Chatbot", page_icon="🤖", layout="centered")
# st.title("🤖 RAG Chatbot")
# st.markdown(
#     """
#     This chatbot answers questions using RAG and LLMs.<br>
#     <small>Choose a model, ask your question, and get instant answers!</small>
#     """,
#     unsafe_allow_html=True,
# )

# if "chat_history" not in st.session_state:
#     st.session_state.chat_history = []

# if "user_id" not in st.session_state:
#     st.session_state.user_id = 1  # For demo, static user id

# st.sidebar.title("Settings")
# model_choice = st.sidebar.radio("Choose Model", ("gemini", "llama"), horizontal=True)

# for entry in st.session_state.chat_history:
#     if entry["role"] == "user":
#         st.chat_message("user").write(entry["content"])
#     else:
#         st.chat_message("assistant").write(entry["content"])

# user_input = st.chat_input("Type your question...")

# if user_input:
#     st.session_state.chat_history.append({"role": "user", "content": user_input})
#     st.chat_message("user").write(user_input)

#     search_engine = st.session_state.search_engine  # Use the cached engine

#     if not search_engine:
#         answer = "Sorry, search system is not ready."
#     else:
#         try:
#             result = process_query(
#                 query=user_input,
#                 min_relevance_score=0.5,
#                 model_choice=model_choice,
#                 gemini_model=gemini_model,
#                 qwen_model=qwen_model,
#                 search_engine=search_engine
#             )
#             answer = result.get("answer", "No answer received.")

#             # --- Improved vague/unsupported answer handling ---
#             fallback_phrases = [
#                 "i cannot answer this question",
#                 "i could not find any relevant information",
#                 "no answer received",
#                 "sorry, an error occurred",
#                 "i cannot answer the question with the provided information",
#                 "i'm sorry, i don't know",
#                 "i don't have enough information",
#                 "i am unable to answer",
#                 "unable to answer",
#                 "no relevant information"
#             ]
#             answer_lower = answer.strip().lower()
#             if any(phrase in answer_lower for phrase in fallback_phrases):
#                 answer = (
#                     "Sorry, I couldn't find an answer to your question."
#                     "Please try rephrasing your question or ask about a specific feature or process."
#                 )
#         except Exception as e:
#             answer = f"Error processing your question: {e}"

#     st.session_state.chat_history.append({"role": "assistant", "content": answer})
#     st.chat_message("assistant").write(answer)





# # import streamlit as st
# # import json
# # import numpy as np
# # from typing import List, Dict, Tuple, Optional
# # from rank_bm25 import BM25Okapi
# # from sentence_transformers import SentenceTransformer
# # from sklearn.metrics.pairwise import cosine_similarity
# # import re
# # import os
# # import google.generativeai as genai
# # from langchain_ollama import ChatOllama
# # import logging
# # import docx
# # import PyPDF2

# # # ----------------------------
# # # Logging Setup
# # # ----------------------------
# # logging.basicConfig(level=logging.INFO)
# # logger = logging.getLogger(__name__)

# # # ----------------------------
# # # Hybrid Search RAG Class (from RAG1.py)
# # # ----------------------------
# # class HybridSearchRAG:
# #     """
# #     Hybrid Search implementation for Retrieval Augmented Generation.
# #     Combines vector-based semantic search with BM25 lexical search.
# #     """

# #     def __init__(
# #             self,
# #             documents: List[Dict],
# #             embedding_model_name: str = "all-MiniLM-L6-v2",
# #             vector_weight: float = 0.7,
# #             bm25_weight: float = 0.3,
# #             top_k: int = 5
# #     ):
# #         self.documents = documents
# #         self.texts = [doc['text'] for doc in documents]
# #         self.doc_ids = [doc.get('id') for doc in documents]
# #         self.top_k = top_k
# #         total = vector_weight + bm25_weight
# #         self.vector_weight = vector_weight / total if total > 0 else 0.5
# #         self.bm25_weight = bm25_weight / total if total > 0 else 0.5

# #         st.info("Loading embedding model...")
# #         self.embedding_model = SentenceTransformer(embedding_model_name)
# #         self.document_embeddings = self._create_document_embeddings()

# #         st.info("Building BM25 index...")
# #         self.bm25_tokenized_corpus = [text.lower().split() for text in self.texts]
# #         self.bm25 = BM25Okapi(self.bm25_tokenized_corpus)
# #         st.success("RAG system is ready!")

# #     def _create_document_embeddings(self) -> np.ndarray:
# #         with st.spinner("Creating document embeddings..."):
# #             embeddings = self.embedding_model.encode(self.texts, show_progress_bar=True)
# #         return embeddings

# #     def _vector_search(self, query: str, top_k: int = None) -> List[Tuple[int, float]]:
# #         if top_k is None:
# #             top_k = self.top_k
# #         query_embedding = self.embedding_model.encode(query)
# #         similarities = cosine_similarity([query_embedding], self.document_embeddings)[0]
# #         top_indices = np.argsort(similarities)[::-1][:top_k]
# #         return [(idx, similarities[idx]) for idx in top_indices]

# #     def _bm25_search(self, query: str, top_k: int = None) -> List[Tuple[int, float]]:
# #         if top_k is None:
# #             top_k = self.top_k
# #         tokenized_query = query.lower().split()
# #         bm25_scores = self.bm25.get_scores(tokenized_query)
# #         top_indices = np.argsort(bm25_scores)[::-1][:top_k]
# #         return [(idx, bm25_scores[idx]) for idx in top_indices]

# #     def search(self, query: str, min_relevance_score: float = 0.5) -> List[Dict]:
# #         vector_results = self._vector_search(query, top_k=self.top_k * 2)
# #         bm25_results = self._bm25_search(query, top_k=self.top_k * 2)

# #         def normalize_scores(results):
# #             scores = [score for _, score in results]
# #             if not scores or max(scores) == min(scores):
# #                 return [(idx, 0.0) for idx, _ in results]
# #             score_range = max(scores) - min(scores)
# #             return [(idx, (score - min(scores)) / score_range if score_range > 0 else 0.0)
# #                     for idx, score in results]

# #         vector_results = normalize_scores(vector_results)
# #         bm25_results = normalize_scores(bm25_results)

# #         combined_scores = {}
# #         for idx, score in vector_results:
# #             combined_scores[idx] = self.vector_weight * score
# #         for idx, score in bm25_results:
# #             if idx in combined_scores:
# #                 combined_scores[idx] += self.bm25_weight * score
# #             else:
# #                 combined_scores[idx] = self.bm25_weight * score

# #         top_results = sorted(combined_scores.items(), key=lambda x: x[1], reverse=True)[:self.top_k]
        
# #         results = []
# #         for idx, score in top_results:
# #             if score >= min_relevance_score:
# #                 doc = self.documents[idx].copy()
# #                 doc['relevance_score'] = float(score)
# #                 results.append(doc)
        
# #         return results

# # # ----------------------------
# # # Query Processing Logic (from webcamaccess.py)
# # # ----------------------------
# # GREETING_PATTERNS = [
# #     r"^h+i+\b", r"^h+e+y+\b", r"^h+e+l+o+\w*\b", r"^good\s*morning\b", r"^good\s*afternoon\b", r"^good\s*evening\b"
# # ]
# # CLOSING_PATTERNS = [
# #     r"thank\s*you\b", r"thanks\b", r"got\s*it\b", r"understood\b", r"bye+\b", r"goodbye\b"
# # ]

# # def is_greeting_only(query: str) -> bool:
# #     q = query.lower().strip()
# #     return any(re.fullmatch(pattern, q) for pattern in GREETING_PATTERNS)

# # def is_closing_only(query: str) -> bool:
# #     q = query.lower().strip()
# #     return any(re.search(pattern, q) for pattern in CLOSING_PATTERNS)

# # def query_starts_with_greeting(query: str) -> bool:
# #     q = query.lower().strip()
# #     return any(re.match(pattern, q) for pattern in GREETING_PATTERNS)

# # def process_query(
# #         query: str,
# #         search_engine: HybridSearchRAG,
# #         model_choice: str,
# #         gemini_model,
# #         ollama_model,
# #         min_relevance_score: float = 0.5
# # ) -> Dict:
    
# #     if is_greeting_only(query):
# #         return {"answer": "Hello! How can I assist you today?"}

# #     if is_closing_only(query):
# #         return {"answer": "You're welcome! Feel free to ask if you have more questions. Goodbye!"}

# #     results = search_engine.search(query, min_relevance_score=min_relevance_score)
    
# #     if not results:
# #         return {"answer": "I could not find any relevant information to answer your question. Please try rephrasing."}

# #     combined_context = "\n\n".join([result['text'] for result in results])
    
# #     prompt = (
# #         "You are a professional customer support assistant.\n"
# #         "Answer the user's question based ONLY on the provided Document Context.\n"
# #         "Do NOT use any external knowledge or make assumptions.\n"
# #         "Keep your response concise, professional, and directly relevant to the user's question.\n"
# #         "If the answer cannot be found in the Document Context, state that you cannot answer the question with the provided information.\n"
# #         "Do not mention the Document Context in your response.\n\n"
# #         f"Document Context:\n{combined_context}\n\n"
# #         f"User Question: {query}"
# #     )

# #     answer = ""
# #     try:
# #         if model_choice == "Gemini":
# #             if gemini_model:
# #                 response = gemini_model.generate_content(prompt)
# #                 answer = response.text.strip()
# #             else:
# #                 answer = "Gemini model is not configured. Please set the API key."
# #         elif model_choice == "Ollama (Local)":
# #             if ollama_model:
# #                 response = ollama_model.invoke(prompt)
# #                 if hasattr(response, "content"):
# #                     answer = response.content.strip()
# #                 else:
# #                     answer = str(response).strip()
# #             else:
# #                 answer = "Ollama model is not available. Please ensure Ollama is running."
# #     except Exception as e:
# #         logger.error(f"Error generating content from {model_choice}: {e}")
# #         answer = f"Sorry, an error occurred while generating the answer with {model_choice}."
    
# #     if query_starts_with_greeting(query) and not is_greeting_only(query):
# #         if not answer.lower().startswith("hello"):
# #              answer = f"Hello! {answer}"
             
# #     return {"answer": answer, "retrieved_context": combined_context}

# # def convert_file_to_documents(uploaded_file, ext):
# #     """Convert uploaded file to a list of {'id': ..., 'text': ...} dicts."""
# #     if ext == ".json":
# #         docs = json.load(uploaded_file)
# #         if isinstance(docs, dict):
# #             docs = [docs]
# #         return docs
# #     elif ext == ".txt":
# #         text = uploaded_file.read().decode("utf-8")
# #         return [{"id": f"txt-{i}", "text": chunk.strip()} for i, chunk in enumerate(text.split("\n\n")) if chunk.strip()]
# #     elif ext == ".csv":
# #         import csv
# #         text = uploaded_file.read().decode("utf-8")
# #         reader = csv.reader(text.splitlines())
# #         rows = [" ".join(row) for row in reader]
# #         return [{"id": f"csv-{i}", "text": row.strip()} for i, row in enumerate(rows) if row.strip()]
# #     elif ext == ".pdf":
# #         pdf_reader = PyPDF2.PdfReader(uploaded_file)
# #         text = ""
# #         for page in pdf_reader.pages:
# #             text += page.extract_text() or ""
# #         return [{"id": f"pdf-{i}", "text": chunk.strip()} for i, chunk in enumerate(text.split("\n\n")) if chunk.strip()]
# #     elif ext == ".docx":
# #         doc = docx.Document(uploaded_file)
# #         text = "\n".join([para.text for para in doc.paragraphs])
# #         return [{"id": f"docx-{i}", "text": chunk.strip()} for i, chunk in enumerate(text.split("\n\n")) if chunk.strip()]
# #     else:
# #         return []

# # # ----------------------------
# # # Streamlit App UI
# # # ----------------------------
# # st.set_page_config(page_title="RAG Chatbot", layout="wide")

# # st.title("📄 Professional RAG Chatbot")
# # st.markdown("This chatbot uses a Retrieval-Augmented Generation model to answer your questions based on a provided knowledge base.")

# # # --- Sidebar for Configuration ---
# # with st.sidebar:
# #     st.header("Configuration")
    
# #     # Model Selection
# #     model_choice = st.selectbox("Choose a Language Model", ["Gemini", "Ollama (Local)"])

# #     # Gemini API Key
# #     if model_choice == "Gemini":
# #         gemini_api_key = st.text_input("Gemini API Key", type="password", help="Get your key from https://aistudio.google.com/app/apikey")
# #         if gemini_api_key:
# #             try:
# #                 genai.configure(api_key=gemini_api_key)
# #             except Exception as e:
# #                 st.error(f"Failed to configure Gemini: {e}")

# #     # Ollama Model Name
# #     if model_choice == "Ollama (Local)":
# #         ollama_model_name = st.text_input("Ollama Model Name", value="phi3:3.8b", help="Ensure the Ollama server is running and the model is pulled.")

# #     # Knowledge Base Upload (now supports more types)
# #     uploaded_file = st.file_uploader(
# #         "Upload Knowledge Base (JSON, TXT, CSV, PDF, DOCX)",
# #         type=["json", "txt", "csv", "pdf", "docx"]
# #     )
# #     st.markdown("""
# #     **Supported formats:** JSON, TXT, CSV, PDF, DOCX  
# #     JSON Example:
# #     ```json
# #     [
# #       {"id": "doc1", "text": "The first document's content..."},
# #       {"id": "doc2", "text": "The second document's content..."}
# #     ]
# #     ```
# #     """)

# # # --- Initialize Models and RAG ---
# # @st.cache_resource
# # def get_gemini_model():
# #     if 'gemini_api_key' in st.session_state and st.session_state.gemini_api_key:
# #         try:
# #             return genai.GenerativeModel("gemini-2.0-flash")
# #         except Exception as e:
# #             st.error(f"Could not initialize Gemini model: {e}")
# #             return None
# #     return None

# # @st.cache_resource
# # def get_ollama_model(_model_name):
# #     try:
# #         return ChatOllama(model=_model_name, temperature=0.1, max_tokens=256)
# #     except Exception as e:
# #         st.error(f"Could not connect to Ollama model '{_model_name}'. Is Ollama running? Error: {e}")
# #         return None

# # @st.cache_resource
# # def get_search_engine(_documents):
# #     if _documents:
# #         return HybridSearchRAG(documents=_documents)
# #     return None

# # gemini_model = None
# # if model_choice == "Gemini" and gemini_api_key:
# #     st.session_state.gemini_api_key = gemini_api_key
# #     gemini_model = get_gemini_model()

# # ollama_model = None
# # if model_choice == "Ollama (Local)":
# #     ollama_model = get_ollama_model(ollama_model_name)

# # search_engine = None
# # if uploaded_file:
# #     ext = os.path.splitext(uploaded_file.name)[1].lower()
# #     try:
# #         documents = convert_file_to_documents(uploaded_file, ext)
# #         if isinstance(documents, list) and all('text' in doc for doc in documents):
# #             search_engine = get_search_engine(documents)
# #         else:
# #             st.error("Invalid file format. Please ensure the file contains text content.")
# #     except Exception as e:
# #         st.error(f"Could not process the uploaded file: {e}")
# # else:
# #     st.warning("Please upload a knowledge base file to begin.")


# # # --- Chat Interface ---
# # if "messages" not in st.session_state:
# #     st.session_state.messages = []

# # for message in st.session_state.messages:
# #     with st.chat_message(message["role"]):
# #         st.markdown(message["content"])
# #         if "context" in message and message["context"]:
# #             with st.expander("Retrieved Context"):
# #                 st.markdown(message["context"])

# # if prompt := st.chat_input("Ask a question..."):
# #     st.session_state.messages.append({"role": "user", "content": prompt})
# #     with st.chat_message("user"):
# #         st.markdown(prompt)

# #     if search_engine:
# #         with st.chat_message("assistant"):
# #             with st.spinner("Thinking..."):
# #                 response = process_query(
# #                     query=prompt,
# #                     search_engine=search_engine,
# #                     model_choice=model_choice,
# #                     gemini_model=gemini_model,
# #                     ollama_model=ollama_model
# #                 )
# #                 answer = response.get("answer", "No answer found.")
# #                 retrieved_context = response.get("retrieved_context")
                
# #                 st.markdown(answer)
                
# #                 if retrieved_context:
# #                     with st.expander("Retrieved Context"):
# #                         st.markdown(retrieved_context)

# #                 st.session_state.messages.append({"role": "assistant", "content": answer, "context": retrieved_context})
# #     else:
# #         st.error("The RAG system is not initialized. Please upload a knowledge base.")
import streamlit as st
import os
import json
import csv
import logging
import requests
from typing import List
from threading import Lock

from RAG1 import HybridSearchRAG
from webcamaccess import process_query
import google.generativeai as genai
from langchain_ollama import ChatOllama
from fileupdater import convert_paragraph_to_json
import docx
import PyPDF2

# --- Logging ---
logging.basicConfig(level=logging.WARNING)
logger = logging.getLogger(__name__)

# --- Gemini Setup ---
genai.configure(api_key=os.getenv("GEMINI_API_KEY", "AIzaSyDxYogQnjYmTbvghsqSWCKtot5DpygC8hA"))

# --- Available Models ---
GEMINI_MODELS = ["gemini-2.0-flash", "gemini-2.0-pro"]

# Function to fetch locally available Ollama models
def get_ollama_models():
    try:
        response = requests.get("http://localhost:11434/api/tags")
        response.raise_for_status()
        models = [model["name"] for model in response.json()["models"]]
        return models if models else ["No models found"]
    except requests.RequestException as e:
        st.sidebar.error(f"Failed to fetch Ollama models: {e}. Ensure Ollama server is running at localhost:11434.")
        return ["phi3:3.8b"]  # Fallback to a default model if fetch fails

OLLAMA_MODELS = get_ollama_models()

# --- Model Initialization ---
def get_gemini_model(model_name):
    try:
        return genai.GenerativeModel(model_name)
    except Exception as e:
        st.sidebar.error(f"Failed to initialize Gemini model {model_name}: {e}")
        return None

def get_ollama_model(model_name):
    try:
        return ChatOllama(model=model_name, temperature=0.1, max_tokens=128)
    except Exception as e:
        st.sidebar.error(f"Failed to initialize Ollama model {model_name}: {e}")
        return None

# --- Document Loading ---
def load_documents(file_paths: List[str]) -> List[dict]:
    documents = []
    for path in file_paths:
        try:
            if path.endswith(".json"):
                with open(path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    if isinstance(data, list):
                        documents.extend(data)
                    else:
                        documents.append(data)
            elif path.endswith(".txt"):
                with open(path, 'r', encoding='utf-8') as f:
                    documents.append({"text": f.read(), "metadata": {}})
            elif path.endswith(".csv"):
                with open(path, 'r', encoding='utf-8') as f:
                    reader = csv.reader(f)
                    text = "\n".join([", ".join(row) for row in reader])
                    documents.append({"text": text, "metadata": {}})
        except Exception as e:
            logger.error(f"Error reading file {path}: {str(e)}")
    return documents

# --- Remove default document path ---
DOCUMENT_PATHS = []  # Start empty, do not load doc1.json

search_engine_lock = Lock()

@st.cache_resource(show_spinner=True)
def get_search_engine(vector_weight, bm25_weight, top_k):
    if not DOCUMENT_PATHS:
        return None
    documents = load_documents(DOCUMENT_PATHS)
    if not documents:
        st.error("No documents loaded. Please upload a knowledge base file.")
        return None
    return HybridSearchRAG(
        documents=documents,
        vector_weight=vector_weight,
        bm25_weight=bm25_weight,
        top_k=top_k
    )

# Only set search_engine after upload
if "search_engine" not in st.session_state:
    st.session_state.search_engine = None
if "topic" not in st.session_state:
    st.session_state.topic = "Unknown Topic"

# --- RAG File Upload Feature ---
st.sidebar.markdown("---")
st.sidebar.subheader("Update Knowledge Base")
uploaded_file = st.sidebar.file_uploader(
    "Upload new RAG file (.json, .txt, .csv, .pdf, .docx)", 
    type=["json", "txt", "csv", "pdf", "docx"],
    help="Replace the current knowledge base with a new file."
)

# Add a button to confirm update
if uploaded_file:
    st.sidebar.success(f"Uploaded and saved as uploaded_files/{uploaded_file.name}")
    if st.sidebar.button("Update Knowledge Base"):
        os.makedirs("uploaded_files", exist_ok=True)
        raw_path = os.path.join("uploaded_files", uploaded_file.name)
        with open(raw_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        ext = os.path.splitext(uploaded_file.name)[1].lower()
        if ext == ".json":
            json_path = raw_path
            try:
                with open(json_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                topic = data.get("topic", "Unknown Topic") if isinstance(data, dict) else "Unknown Topic"
            except Exception:
                topic = "Unknown Topic"
        else:
            uploaded_file.seek(0)
            if ext == ".txt":
                paragraph = uploaded_file.read().decode("utf-8")
            elif ext == ".csv":
                paragraph = uploaded_file.read().decode("utf-8")
            elif ext == ".pdf":
                try:
                    pdf_reader = PyPDF2.PdfReader(uploaded_file)
                    paragraph = ""
                    for page in pdf_reader.pages:
                        paragraph += page.extract_text() or ""
                except Exception as e:
                    st.sidebar.error(f"PDF extraction failed: {e}")
                    paragraph = ""
            elif ext == ".docx":
                try:
                    doc = docx.Document(uploaded_file)
                    paragraph = "\n".join([para.text for para in doc.paragraphs])
                except Exception as e:
                    st.sidebar.error(f"DOCX extraction failed: {e}")
                    paragraph = ""
            else:
                paragraph = ""

            if paragraph.strip():
                result = convert_paragraph_to_json(paragraph)
                topic = result.get("topic", "Unknown Topic")
                new_entries = result.get("entries", [])
                base_name = os.path.splitext(uploaded_file.name)[0]
                json_path = os.path.join("uploaded_files", f"{base_name}.json")
                with open(json_path, "w", encoding="utf-8") as jf:
                    json.dump(new_entries, jf, indent=2, ensure_ascii=False)
                st.sidebar.info(f"Converted and saved as {json_path}")
            else:
                st.sidebar.error("No text could be extracted from the uploaded file.")
                json_path = None
                topic = "Unknown Topic"

        if json_path:
            DOCUMENT_PATHS.clear()
            DOCUMENT_PATHS.append(json_path)
            get_search_engine.clear()
            st.session_state.search_engine = get_search_engine(
                st.session_state.get("vector_weight", 0.7),
                st.session_state.get("bm25_weight", 0.3),
                st.session_state.get("top_k", 3)
            )
            st.session_state.topic = topic
            st.session_state.chat_history = []
            st.session_state.last_uploaded = uploaded_file.name
            st.sidebar.success(f"Knowledge base updated for {topic}! You can now ask questions about the new document.")

# --- Block Chat Until Knowledge Base is Loaded ---
def is_kb_loaded():
    return st.session_state.get("search_engine") is not None

if not is_kb_loaded():
    st.warning("Please upload a knowledge base file (.json, .txt, .csv, .pdf, .docx) to start using the chatbot.")
    st.stop()

# --- Streamlit UI ---
st.set_page_config(page_title="RAG Chatbot", page_icon="🤖", layout="centered")
st.title(f"🤖 RAG Chatbot for {st.session_state.topic}")
st.markdown(
    f"""
    This chatbot answers questions about {st.session_state.topic} using RAG and LLMs.<br>
    <small>Choose a model, ask your question, and get instant answers!</small>
    """,
    unsafe_allow_html=True,
)

if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

if "user_id" not in st.session_state:
    st.session_state.user_id = 1  # For demo, static user id

# --- Sidebar Settings ---
st.sidebar.title("Settings")

# Model Selection Dropdown
model_type = st.sidebar.selectbox("LLM Selection", ["ONLINE (Gemini)", "OFFLINE (Ollama)"])
if model_type == "ONLINE (Gemini)":
    model_name = st.sidebar.selectbox("Gemini Model", GEMINI_MODELS, index=0)
    model = get_gemini_model(model_name) if model_name else None
elif model_type == "OFFLINE (Ollama)":
    model_name = st.sidebar.selectbox("Ollama Model", OLLAMA_MODELS, index=0)
    model = get_ollama_model(model_name) if model_name else None
else:
    model = None
    st.sidebar.error("Please select a valid model type.")

# Search Settings
st.sidebar.subheader("Search Settings")
vector_weight = st.sidebar.slider("Vector Search Weight", 0.0, 1.0, st.session_state.get("vector_weight", 0.7), 0.1)
bm25_weight = st.sidebar.slider("BM25 Weight", 0.0, 1.0, st.session_state.get("bm25_weight", 0.3), 0.1)
top_k = st.sidebar.slider("Number of Results", 1, 10, st.session_state.get("top_k", 3), 1)
min_relevance_score = st.sidebar.slider("Minimum Relevance Score", 0.0, 1.0, st.session_state.get("min_relevance_score", 0.5), 0.1)

# Store search settings in session state
st.session_state.vector_weight = vector_weight
st.session_state.bm25_weight = bm25_weight
st.session_state.top_k = top_k
st.session_state.min_relevance_score = min_relevance_score

# Update search engine if settings change
if "search_engine" in st.session_state and (st.session_state.get("vector_weight") != vector_weight or
                                           st.session_state.get("bm25_weight") != bm25_weight or
                                           st.session_state.get("top_k") != top_k):
    get_search_engine.clear()
    st.session_state.search_engine = get_search_engine(vector_weight, bm25_weight, top_k)

for entry in st.session_state.chat_history:
    if entry["role"] == "user":
        st.chat_message("user").write(entry["content"])
    else:
        st.chat_message("assistant").write(entry["content"])

user_input = st.chat_input(f"Type your question about {st.session_state.topic}...")

if user_input:
    st.session_state.chat_history.append({"role": "user", "content": user_input})
    st.chat_message("user").write(user_input)

    search_engine = st.session_state.search_engine

    if not search_engine or not model:
        answer = "Sorry, search system or model is not ready."
    else:
        try:
            result = process_query(
                query=user_input,
                min_relevance_score=min_relevance_score,
                model_choice="gemini" if model_type == "ONLINE (Gemini)" else "llama",
                gemini_model=model if model_type == "ONLINE (Gemini)" else None,
                qwen_model=model if model_type == "OFFLINE (Ollama)" else None,
                search_engine=search_engine,
                topic=st.session_state.topic
            )
            answer = result.get("answer", "No answer received.")

            # Improved vague/unsupported answer handling
            fallback_phrases = [
                "i cannot answer this question",
                "i could not find any relevant information",
                "no answer received",
                "sorry, an error occurred",
                "i cannot answer the question with the provided information",
                "i'm sorry, i don't know",
                "i don't have enough information",
                "i am unable to answer",
                "unable to answer",
                "no relevant information"
            ]
            answer_lower = answer.strip().lower()
            if any(phrase in answer_lower for phrase in fallback_phrases):
                answer = (
                    f"Sorry, I couldn't find an answer to your question about {st.session_state.topic}."
                    "Please try rephrasing your question or ask about a specific feature or process."
                )
        except Exception as e:
            answer = f"Error processing your question: {e}"

    st.session_state.chat_history.append({"role": "assistant", "content": answer})
    st.chat_message("assistant").write(answer)